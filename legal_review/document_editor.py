import io
import datetime
import html
import re
import streamlit as st
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

try:
    import docx
except ImportError:
    pass

from legal_review.text_matcher import find_best_text_span


def _apply_text_modifications(text: str, modifications: list[dict]) -> str:
    final_text = text or ""
    for mod in modifications:
        original = (mod.get("original") or "").strip()
        suggestion = (mod.get("suggestion") or "").strip()
        if original and suggestion:
            final_text = final_text.replace(original, suggestion)
    return final_text


def _get_export_risks(snapshot: dict | None) -> list[dict]:
    snap = snapshot or {}
    return snap.get("raw_export_risks") or snap.get("risks") or []


def _build_filtered_mapping(full_text: str, keep_char) -> tuple[str, list[int]]:
    chars: list[str] = []
    mapping: list[int] = []
    for idx, ch in enumerate(full_text or ""):
        if keep_char(ch):
            chars.append(ch)
            mapping.append(idx)
    return "".join(chars), mapping


def _find_strict_span_in_paragraph(paragraph_text: str, query: str) -> tuple[int, int] | None:
    source = paragraph_text or ""
    target = (query or "").strip()
    if not source or not target:
        return None

    pos = source.find(target)
    if pos != -1:
        return pos, pos + len(target)

    filtered_source, source_map = _build_filtered_mapping(source, lambda ch: not ch.isspace())
    filtered_target, _ = _build_filtered_mapping(target, lambda ch: not ch.isspace())
    if filtered_target:
        pos = filtered_source.find(filtered_target)
        if pos != -1:
            return source_map[pos], source_map[pos + len(filtered_target) - 1] + 1

    filtered_source_punc, source_map_punc = _build_filtered_mapping(source, lambda ch: bool(re.match(r"[\w\u4e00-\u9fa5]", ch)))
    filtered_target_punc, _ = _build_filtered_mapping(target, lambda ch: bool(re.match(r"[\w\u4e00-\u9fa5]", ch)))
    if filtered_target_punc:
        pos = filtered_source_punc.find(filtered_target_punc)
        if pos != -1:
            return source_map_punc[pos], source_map_punc[pos + len(filtered_target_punc) - 1] + 1

    return None


def _replace_paragraph_text(paragraph, start: int, end: int, replacement: str) -> None:
    para_text = paragraph.text or ""
    if start < 0 or end < start or end > len(para_text):
        return

    if not paragraph.runs:
        paragraph.add_run(para_text[:start] + replacement + para_text[end:])
        return

    cursor = 0
    inserted = False
    for run in paragraph.runs:
        run_text = run.text or ""
        run_start = cursor
        run_end = cursor + len(run_text)
        cursor = run_end

        if run_end <= start or run_start >= end:
            continue

        local_start = max(0, start - run_start)
        local_end = min(len(run_text), end - run_start)
        prefix = run_text[:local_start]
        suffix = run_text[local_end:]

        if not inserted:
            run.text = prefix + replacement + suffix
            inserted = True
        else:
            run.text = prefix + suffix

    if not inserted:
        paragraph.runs[0].text = para_text[:start] + replacement + para_text[end:]


def _get_runs_for_span(paragraph, start: int, end: int) -> list:
    runs = list(paragraph.runs)
    if not runs:
        runs = [paragraph.add_run("")]
    if end <= start:
        end = start + 1
    cursor = 0
    selected = []
    for run in runs:
        run_text = run.text or ""
        run_start = cursor
        run_end = cursor + len(run_text)
        cursor = run_end
        if run_end <= start or run_start >= end:
            continue
        selected.append(run)
    return selected or runs[:1]


def _ensure_anchor_run(paragraph):
    if paragraph.runs:
        return paragraph.runs[0]
    return paragraph.add_run("")


def _add_comment(doc, runs, text: str) -> None:
    try:
        doc.add_comment(runs, text, author="智审法务", initials="AI")
    except Exception:
        pass


def _highlight_runs(runs: list, color=WD_COLOR_INDEX.YELLOW) -> None:
    for run in runs or []:
        try:
            run.font.highlight_color = color
        except Exception:
            pass


def _insert_manual_review_marker_before(paragraph, original: str, suggestion: str) -> None:
    try:
        marker_para = paragraph.insert_paragraph_before()
    except Exception:
        return
    title = marker_para.add_run("【人工核查：本处未自动修改】")
    title.bold = True
    title.font.highlight_color = WD_COLOR_INDEX.YELLOW
    title.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    original_preview = (original or "").strip()
    suggestion_preview = (suggestion or "").strip()
    if len(original_preview) > 80:
        original_preview = original_preview[:80] + "..."
    if len(suggestion_preview) > 80:
        suggestion_preview = suggestion_preview[:80] + "..."

    body = marker_para.add_run(
        f" 原文：{original_preview}；建议：{suggestion_preview}"
    )
    body.font.highlight_color = WD_COLOR_INDEX.YELLOW
    body.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _build_docx_paragraph_spans(paragraphs) -> list[dict]:
    spans: list[dict] = []
    cursor = 0
    for idx, para in enumerate(paragraphs):
        text = para.text or ""
        start = cursor
        end = start + len(text)
        spans.append({"index": idx, "paragraph": para, "start": start, "end": end})
        cursor = end
        if idx < len(paragraphs) - 1:
            cursor += 1
    return spans


def _build_document_text_from_paragraph_spans(paragraph_spans: list[dict]) -> str:
    return "\n".join(item["paragraph"].text or "" for item in paragraph_spans)


def _find_paragraph_range_for_span(paragraph_spans: list[dict], start: int, end: int) -> tuple[int, int] | None:
    first = None
    last = None
    for idx, item in enumerate(paragraph_spans):
        para_start = item["start"]
        para_end = item["end"]
        overlaps = not (end <= para_start or start >= para_end)
        touches_empty_para = para_start == para_end and start == end == para_start
        if overlaps or touches_empty_para:
            if first is None:
                first = idx
            last = idx
    if first is None or last is None:
        return None
    return first, last


def _apply_span_replacement_to_docx(paragraph_spans: list[dict], start: int, end: int, replacement: str) -> bool:
    target_range = _find_paragraph_range_for_span(paragraph_spans, start, end)
    if target_range is None:
        return False

    first_idx, last_idx = target_range
    first_item = paragraph_spans[first_idx]
    last_item = paragraph_spans[last_idx]
    block_start = first_item["start"]
    block_end = last_item["end"]
    block_text = "\n".join(item["paragraph"].text or "" for item in paragraph_spans[first_idx : last_idx + 1])
    rel_start = start - block_start
    rel_end = end - block_start
    if rel_start < 0 or rel_end < rel_start or rel_end > len(block_text):
        return False

    if first_idx == last_idx:
        _replace_paragraph_text(first_item["paragraph"], rel_start, rel_end, replacement)
        return True

    new_block_text = block_text[:rel_start] + replacement + block_text[rel_end:]
    new_lines = new_block_text.split("\n")
    affected = paragraph_spans[first_idx : last_idx + 1]
    for idx, item in enumerate(affected):
        _set_paragraph_text(item["paragraph"], new_lines[idx] if idx < len(new_lines) else "")
    if len(new_lines) > len(affected):
        overflow = "\n".join(new_lines[len(affected) - 1 :])
        _set_paragraph_text(affected[-1]["paragraph"], overflow)
    return True


def _find_docx_span_by_query(paragraph_spans: list[dict], query: str) -> tuple[int, int] | None:
    full_text = _build_document_text_from_paragraph_spans(paragraph_spans)
    return _find_strict_span_in_paragraph(full_text, query)


def _find_docx_anchor_by_query(paragraph_spans: list[dict], query: str) -> tuple[int, int] | None:
    full_text = _build_document_text_from_paragraph_spans(paragraph_spans)
    strict_match = _find_strict_span_in_paragraph(full_text, query)
    if strict_match is not None:
        return strict_match
    start, end = find_best_text_span(full_text, query)
    if start == -1 or end == -1:
        return None
    return start, end


def _append_manual_review_block(doc, unresolved_entries: list[dict]) -> None:
    if not unresolved_entries:
        return
    heading = doc.add_paragraph()
    run = heading.add_run("【人工核查清单】以下修改未能自动回写，请按批注或下列清单手动检查")
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for idx, entry in enumerate(unresolved_entries, start=1):
        para = doc.add_paragraph()
        title = para.add_run(f"{idx}. 未自动替换\n")
        title.bold = True
        title.font.highlight_color = WD_COLOR_INDEX.YELLOW
        para.add_run(f"原文：{entry.get('original', '')}\n")
        para.add_run(f"建议：{entry.get('suggestion', '')}\n")
        para.add_run(f"原因：{entry.get('reason', '请人工核查')}")


def collect_applied_modifications(snapshot: dict | None, applied_indices: set[int]) -> list[tuple[int, dict]]:
    risks = _get_export_risks(snapshot)
    collected: list[tuple[int, dict]] = []
    for idx in sorted(applied_indices or set()):
        if 0 <= idx < len(risks):
            collected.append((idx, risks[idx]))
    return collected


def render_export_change_summary(modifications: list[tuple[int, dict]]) -> None:
    if not modifications:
        st.info("⚠️ 您尚未在工作台中采纳任何修改。当前导出的文件将与原文一致。")
        return

    st.markdown("#### 本次纳入导出的变更")
    for idx, risk in modifications:
        original = (risk.get("original") or "").strip()
        suggestion = (risk.get("suggestion_display") or risk.get("suggestion") or "").strip()
        export_warning = (risk.get("export_warning") or "").strip()
        warning_html = ""
        if export_warning:
            warning_html = (
                '<div style="font-family:Outfit,Noto Sans SC,sans-serif;'
                'font-size:0.82rem;color:#b15b52;line-height:1.7;margin-top:6px;">'
                f'{html.escape(export_warning)}</div>'
            )
        st.markdown(
            f'<div style="padding:14px 16px;border-radius:16px;border:1px solid #e2ddd5;background:#fffdf9;margin:10px 0;">'
            f'<div style="font-family:Outfit,Noto Sans SC,sans-serif;font-size:0.8rem;color:#8a857f;margin-bottom:6px;">风险点 {idx + 1}</div>'
            f'<div style="font-family:Outfit,Noto Sans SC,sans-serif;font-size:0.88rem;color:#1a1f2e;line-height:1.7;"><strong>原文：</strong>{html.escape(original[:160])}</div>'
            f'<div style="font-family:Outfit,Noto Sans SC,sans-serif;font-size:0.88rem;color:#1a1f2e;line-height:1.7;margin-top:6px;"><strong>替换后：</strong>{html.escape(suggestion[:180])}</div>'
            f'{warning_html}'
            f'</div>',
            unsafe_allow_html=True,
        )

def apply_changes_to_docx(original_bytes: bytes, modifications: list) -> tuple[io.BytesIO, list[dict]]:
    """
    修改内存中的 docx 文件并在原格式下返回新的二进制流。
    """
    doc = docx.Document(io.BytesIO(original_bytes))

    all_paras = list(doc.paragraphs)
    paragraph_spans = _build_docx_paragraph_spans(all_paras)

    unresolved_entries: list[dict] = []
    ordered_modifications = sorted(
        modifications,
        key=lambda mod: int(((mod or {}).get("raw_original_span") or {}).get("start", -1)),
        reverse=True,
    )
    for mod in ordered_modifications:
        orig_text = (mod.get("original") or "").strip()
        sugg_text = (mod.get("suggestion") or "").strip()
        if not orig_text or not sugg_text:
            continue
        raw_span = (mod.get("raw_original_span") or {})
        span_start = raw_span.get("start")
        span_end = raw_span.get("end")

        replaced = False
        if isinstance(span_start, int) and isinstance(span_end, int):
            replaced = _apply_span_replacement_to_docx(paragraph_spans, span_start, span_end, sugg_text)
            if replaced:
                target_range = _find_paragraph_range_for_span(paragraph_spans, span_start, span_end)
                if target_range is not None:
                    first_idx, last_idx = target_range
                    if first_idx == last_idx:
                        para = paragraph_spans[first_idx]["paragraph"]
                        comment_runs = _get_runs_for_span(para, span_start - paragraph_spans[first_idx]["start"], span_start - paragraph_spans[first_idx]["start"] + len(sugg_text))
                    else:
                        first_para = paragraph_spans[first_idx]["paragraph"]
                        last_para = paragraph_spans[last_idx]["paragraph"]
                        comment_runs = [_ensure_anchor_run(first_para), _ensure_anchor_run(last_para)]
                    _add_comment(
                        doc,
                        comment_runs,
                        f"已自动应用修改。\n原文：{orig_text}\n修改后：{sugg_text}",
                    )

        if not replaced:
            query_span = _find_docx_span_by_query(paragraph_spans, orig_text)
            if query_span is not None:
                replaced = _apply_span_replacement_to_docx(paragraph_spans, query_span[0], query_span[1], sugg_text)
                if replaced:
                    target_range = _find_paragraph_range_for_span(paragraph_spans, query_span[0], query_span[1])
                    if target_range is not None:
                        first_idx, last_idx = target_range
                        if first_idx == last_idx:
                            para = paragraph_spans[first_idx]["paragraph"]
                            comment_runs = _get_runs_for_span(para, query_span[0] - paragraph_spans[first_idx]["start"], query_span[0] - paragraph_spans[first_idx]["start"] + len(sugg_text))
                        else:
                            first_para = paragraph_spans[first_idx]["paragraph"]
                            last_para = paragraph_spans[last_idx]["paragraph"]
                            comment_runs = [_ensure_anchor_run(first_para), _ensure_anchor_run(last_para)]
                        _add_comment(
                            doc,
                            comment_runs,
                            f"已自动应用修改。\n原文：{orig_text}\n修改后：{sugg_text}",
                        )

        if not replaced:
            candidates: list[tuple[object, tuple[int, int]]] = []
            for para in all_paras:
                span = _find_strict_span_in_paragraph(para.text or "", orig_text)
                if span is not None:
                    candidates.append((para, span))
            if len(candidates) == 1:
                target_para, (start_in_para, end_in_para) = candidates[0]
                _replace_paragraph_text(target_para, start_in_para, end_in_para, sugg_text)
                comment_runs = _get_runs_for_span(target_para, start_in_para, start_in_para + len(sugg_text))
                _add_comment(
                    doc,
                    comment_runs,
                    f"已自动应用修改。\n原文：{orig_text}\n修改后：{sugg_text}",
                )
                replaced = True

        if not replaced:
            entry = {
                "original": orig_text,
                "suggestion": sugg_text,
                "reason": "系统未能将该修改精确回写到原文，请人工核查。",
            }
            query_span = _find_docx_anchor_by_query(paragraph_spans, orig_text)
            if query_span is not None:
                target_range = _find_paragraph_range_for_span(paragraph_spans, query_span[0], query_span[1])
                if target_range is not None:
                    first_idx, last_idx = target_range
                    first_para = paragraph_spans[first_idx]["paragraph"]
                    last_para = paragraph_spans[last_idx]["paragraph"]
                    if first_idx == last_idx:
                        runs = _get_runs_for_span(first_para, query_span[0] - paragraph_spans[first_idx]["start"], query_span[1] - paragraph_spans[first_idx]["start"])
                    else:
                        runs = [_ensure_anchor_run(first_para), _ensure_anchor_run(last_para)]
                    _highlight_runs(runs)
                    _insert_manual_review_marker_before(first_para, orig_text, sugg_text)
                    _add_comment(
                        doc,
                        runs,
                        f"未自动应用修改，请人工核查。\n原文：{orig_text}\n建议：{sugg_text}",
                    )
            unresolved_entries.append(entry)

    _append_manual_review_block(doc, unresolved_entries)
    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream, unresolved_entries

def execute_export_pipeline():
    """
    在主界面中渲染导出操作面板。
    """
    st.markdown("### 导出与变更确认")

    snap = st.session_state.get("review_snapshot")
    applied_idx = st.session_state.get("applied_risks", set())

    if not snap or not snap.get("risks"):
        st.error("没有可供导出的审查数据。")
        if st.button("关闭面板"):
            st.session_state["show_export_dialog"] = False
            st.rerun()
        return

    modifications_with_idx = collect_applied_modifications(snap, applied_idx)
    modifications = [risk for _, risk in modifications_with_idx]
    render_export_change_summary(modifications_with_idx)

    if modifications:
        st.success(f"准备导出：已经包含 {len(modifications)} 处您确认应用的专业修改。")

    masking = snap.get("masking") or {}
    unresolved_candidates = [
        (idx, risk) for idx, risk in modifications_with_idx
        if not bool(risk.get("export_actionable", True))
    ]

    if masking.get("sent_to_model_masked"):
        if modifications:
            st.success("本次审查虽使用脱敏文本，但系统已将可定位条款映射回原文，可继续导出正式版本。")
        else:
            st.warning("本次审查使用脱敏文本，但当前已采纳条款未能可靠映射回原文，暂时无法生成正式修订稿。")

    if unresolved_candidates:
        candidate_labels = "、".join(f"风险点 {idx + 1}" for idx, _ in unresolved_candidates[:6])
        if len(unresolved_candidates) > 6:
            candidate_labels += " 等"
        st.info(f"{candidate_labels} 可能需要人工复核；导出 Word 时系统会继续尝试回写，若仍失败会在文档中用高亮和批注标出。")

    orig_bytes = st.session_state.get("original_file_bytes")
    orig_name = st.session_state.get("original_file_name") or "contract"
    
    now_str = datetime.datetime.now().strftime('%Y%m%d_%H%M')
    is_docx = orig_name.lower().endswith(".docx")

    if not modifications:
        final_text = _apply_text_modifications(snap.get("raw_text") or snap.get("review_text") or snap.get("text", ""), modifications)
        col_a, col_b = st.columns(2)
        with col_a:
            st.download_button(
                "⬇️ 下载当前正文 (.txt)",
                data=final_text.encode("utf-8"),
                file_name=f"合同正文_{now_str}.txt",
                mime="text/plain",
                use_container_width=True
            )
        with col_b:
            if st.button("关闭面板", use_container_width=True):
                st.session_state["show_export_dialog"] = False
                st.rerun()
        st.markdown("---")
        return
    
    if orig_bytes and is_docx:
        st.success("✅ 检测到了原始的 Word (.docx) 文档格式！我们将执行格式无损级别的底层内容替换。")
        try:
            out_bytesIO, unresolved_entries = apply_changes_to_docx(orig_bytes, modifications)
            if unresolved_entries:
                preview = "；".join((entry.get("original", "")[:24] + ("..." if len(entry.get("original", "")) > 24 else "")) for entry in unresolved_entries[:3])
                if len(unresolved_entries) > 3:
                    preview += "；……"
                st.warning(f"有 {len(unresolved_entries)} 条修改未能自动回写。导出的 Word 会在正文附近插入人工核查提示，并配合黄色高亮、批注和文末核查清单提示你手动检查：{preview}")
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "⬇️ 点击下载 原格式修改版 (.docx)",
                    data=out_bytesIO.getvalue(),
                    file_name=f"已修订_{orig_name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
            with col2:
                if st.button("关闭面板", use_container_width=True):
                    st.session_state["show_export_dialog"] = False
                    st.rerun()
        except Exception as e:
            st.error(f"处理 docx 文件时内部发生错误: {e}")
            if st.button("关闭面板"):
                st.session_state["show_export_dialog"] = False
                st.rerun()
    else:
        st.warning("⚠️ 由于原始文件不是纯正的 Word (.docx) 或者您是直接复制粘贴的纯文本（比如 PDF 等格式固定排版），程序将通过应用补丁为您打包一份标准的基础文档。")
        final_text = _apply_text_modifications(
            snap.get("raw_text") or snap.get("review_text") or snap.get("text", ""),
            modifications,
        )
                
        # 提供 TXT 与普通生成的 DOCX 下载
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            st.download_button(
                "⬇️ 下载为普通文本 (.txt)",
                data=final_text.encode("utf-8"),
                file_name=f"修订_合同_{now_str}.txt",
                mime="text/plain",
                use_container_width=True
            )
        with col_b:
            try:
                doc = docx.Document()
                for line in final_text.split("\n"):
                    doc.add_paragraph(line)
                bf = io.BytesIO()
                doc.save(bf)
                st.download_button(
                    "⬇️ 下载为普通版Word (.docx)",
                    data=bf.getvalue(),
                    file_name=f"修订_无格式版本_{now_str}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"生成 Word 备选时失败: {e}")
                
        with col_c:
            if st.button("关闭面板", use_container_width=True):
                st.session_state["show_export_dialog"] = False
                st.rerun()

    st.markdown("---")
